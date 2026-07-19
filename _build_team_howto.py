"""Generates the team 'how to' Word doc for v2 (Astons Invoice Platform)."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASTONS_GREEN = RGBColor(0x1a, 0x5c, 0x2e)
ASTONS_MID = RGBColor(0x3a, 0x8c, 0x4e)
GREY = RGBColor(0x55, 0x55, 0x55)

APP_URL = "https://astons-invoice-platform-production.up.railway.app"

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def heading(text, level=1, colour=ASTONS_GREEN):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = colour
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(text, bold=False, italic=False, colour=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(title, body, colour=ASTONS_GREEN):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    # Shade cell
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E8F2EA")
    tc_pr.append(shd)

    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.color.rgb = colour
    r1.font.size = Pt(11)

    p2 = cell.add_paragraph()
    p2.add_run(body).font.size = Pt(11)

    doc.add_paragraph()


# === HEADER / LOGO ===
logo_path = Path(__file__).parent / "astons_logo.png"
if logo_path.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(str(logo_path), width=Inches(1.6))

heading("Astons Invoice Platform", level=1)
para(
    "How to submit invoices for approval — walkthrough for the team",
    italic=True,
    colour=GREY,
)
doc.add_paragraph()

# === WHAT THIS IS ===
heading("What this is", level=2)
para(
    "This is our new web-based invoice tool. It replaces the old "
    "manual process of rebuilding BrightManager fee notes in Word / "
    "PDF editors. You drop the BrightManager PDF in, pick the portfolio, "
    "and it produces a properly branded Astons invoice."
)
para(
    "Every invoice goes through Ash for approval before it can be sent "
    "to the client — so don't worry about getting things wrong, nothing "
    "leaves the platform without review."
)

# === ACCESS ===
heading("How to get in", level=2)
numbered("the link below in any browser (Chrome, Edge, Safari — all fine). "
         "You can bookmark it.", bold_prefix="Open ")
para(APP_URL, bold=True, colour=ASTONS_MID)
numbered("in with the username and temporary password Ash has sent you.",
         bold_prefix="Sign ")
numbered("your password after first login: click the ", bold_prefix="Change ")
para(
    "   (You cannot do this yourself — message Ash and he'll reset it for you "
    "from the admin panel, then you sign in with the new one.)",
    italic=True,
    colour=GREY,
)

callout(
    "Works on any device",
    "Desktop, laptop, tablet, even phone in a pinch. No install needed. "
    "If you get a login screen, you're in the right place.",
)

# === SUBMITTING AN INVOICE ===
heading("How to submit an invoice", level=2)

numbered("Download the fee note PDF from BrightManager as you normally would.",
         bold_prefix="Step 1 — ")
numbered("On the platform, you'll land on the ",
         bold_prefix="Step 2 — ")
para(
    "   \"New submission\" tab by default.",
    colour=GREY,
)
numbered("which portfolio this invoice is for:",
         bold_prefix="Step 3 — Pick ")
bullet("60-83-71 / 19010489", bold_prefix="A-Portfolio — ")
bullet("04-13-76 / 00273335", bold_prefix="C-Portfolio — ")
para(
    "   If you're not sure which portfolio a client belongs to, check "
    "BrightManager or ask Ash before submitting. This decides which "
    "bank account appears on the invoice.",
    italic=True,
    colour=GREY,
)

numbered("or drop the BrightManager PDF(s) onto the upload box. "
         "You can do several at once.", bold_prefix="Step 4 — Drag ")

numbered("\"Generate branded previews\". You'll get a preview of each "
         "branded Astons invoice.", bold_prefix="Step 5 — Click ")

numbered("each preview — open the PDF and check:",
         bold_prefix="Step 6 — Review ")
bullet("Client name is spelt correctly")
bullet("Invoice number matches BrightManager")
bullet("Line items and total look right")
bullet("Bank details at the bottom are the correct portfolio")

numbered("\"Submit for approval\" on each one (or use ",
         bold_prefix="Step 7 — Click ")
para(
    "   \"Submit all remaining\" if you've done a batch).",
    colour=GREY,
)

callout(
    "What 'Submitted' means",
    "Once you click Submit, the invoice goes to Ash's approval queue. "
    "It is NOT sent to the client yet. You'll see a green 'Submitted' "
    "confirmation, and the invoice will now appear under your 'My invoices' tab.",
)

# === AFTER SUBMISSION ===
heading("After you submit", level=2)

para("Click the ", size=11)
para("\"My invoices\" tab to see everything you've submitted. Each invoice will show one of four statuses:",
     size=11)
bullet("waiting for Ash to review.",
       bold_prefix="Pending approval — ")
bullet("Ash has approved it. A \"Download branded PDF\" button now appears.",
       bold_prefix="Approved — ")
bullet("Ash has sent it back. Read the rejection note, fix whatever's wrong, and resubmit.",
       bold_prefix="Rejected — ")
bullet("you've downloaded and emailed it to the client, and marked it as sent.",
       bold_prefix="Sent — ")

# === SENDING TO CLIENT ===
heading("Sending the invoice to the client", level=2)
numbered("for your approved invoice in \"My invoices\".",
         bold_prefix="Wait ")
numbered("\"Download branded PDF\".",
         bold_prefix="Click ")
numbered("the PDF to the client exactly as you normally would "
         "(via Outlook / whatever template you use).",
         bold_prefix="Email ")
numbered("\"Mark as sent\" on the platform so we have a record.",
         bold_prefix="Come back and click ")

para(
    "That's it — you're done with that invoice.",
    bold=True,
    colour=ASTONS_GREEN,
)

# === COMMON QUESTIONS ===
heading("Common questions", level=2)

para("\"I uploaded the wrong PDF / picked the wrong portfolio\"", bold=True)
para(
    "Before submitting: just upload the correct one again, or change the portfolio "
    "— the previews will regenerate.",
)
para(
    "After submitting: don't worry, just message Ash and he'll reject it in the "
    "queue so you can resubmit.",
)

para("\"I got an error when generating the preview\"", bold=True)
para(
    "Usually means something in the BrightManager PDF is unusual. Screenshot "
    "the error and send it to Ash with the original PDF — he'll have a look.",
)

para("\"Can I edit the generated invoice?\"", bold=True)
para(
    "No — the whole point is that the platform generates it consistently. "
    "If something's wrong on the invoice, fix it in BrightManager first "
    "and re-upload.",
)

para("\"What if I need to send an invoice urgently?\"", bold=True)
para(
    "Ping Ash directly so he knows to approve it quickly. Approval is usually "
    "same-day but won't be instant.",
)

# === WHAT WE WANT FROM TESTING ===
heading("What we'd like from you during testing", level=2)
para(
    "Please try submitting a few real invoices through the platform over the next "
    "week and let Ash know:",
)
bullet("Anything that didn't work as expected")
bullet("Anything that felt awkward or slow")
bullet("Anything missing that you'd want it to do")
bullet("Any typos, wrong numbers, or formatting issues on the branded PDFs")

para(
    "Screenshots are really helpful if something goes wrong — even a phone photo "
    "of the screen is fine.",
    italic=True,
    colour=GREY,
)

# === PRIVACY / SAFETY ===
heading("A note on privacy", level=2)
para(
    "Everything on the platform stays inside Astons. It's hosted on our private "
    "Railway deployment and only people with a login can access it. Invoices, "
    "client names, and amounts are never shared with anyone outside the firm."
)

doc.add_paragraph()
para("— Ash", italic=True, colour=GREY)

# === SAVE ===
out_path = Path.home() / "OneDrive - Astons Accountants" / "Desktop" / "Invoice_Platform_Team_HowTo.docx"
out_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out_path))
print(f"Saved: {out_path}")
print(f"Size: {out_path.stat().st_size} bytes")
